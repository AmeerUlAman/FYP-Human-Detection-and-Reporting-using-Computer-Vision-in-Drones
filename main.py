import cv2
import time
import math
import os
from tkinter import *
from tkinter import messagebox, filedialog
import threading
from PIL import Image
from docx import Document
from docx.shared import Inches
from ultralytics import YOLO

#We selected YOLO 
#ThisLOC Load the YOLO model for object detection
model = YOLO("yolov8n.pt")  

# Folder to save images and Word report
output_folder = "detections"
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# To store the last saved bounding box coordinates and time
last_saved_boxes = []
last_save_time = 0
image_counter = 0
save_delay = 0  # Minimum time delay in seconds between saving images
proximity_threshold = 50  # The maximum allowed distance (in pixels) to consider bounding boxes as the same

# Function to generate a unique filename based on the current timestamp
def generate_unique_filename():
    timestamp = time.strftime("%Y%m%d-%H%M%S")  # Get current timestamp for unique naming
    return os.path.join(output_folder, f"human_detections_report_{timestamp}.docx")

# Initialize a list to store human IDs and image file paths
human_images = []

# Function to calculate Euclidean distance between two bounding boxes (centers)
def calculate_distance(box1, box2):
 
    center1 = ((box1[0] + box1[2]) / 2, (box1[1] + box1[3]) / 2)
    center2 = ((box2[0] + box2[2]) / 2, (box2[1] + box2[3]) / 2)
 
    distance = math.sqrt((center1[0] - center2[0]) ** 2 + (center1[1] - center2[1]) ** 2)
    return distance

# Function to create the Word document report
def create_doc_report(doc, total_humans):
    doc.add_heading('Human Detection Report', 0)
    doc.add_paragraph(f'Total Detected Humans: {total_humans}')
    doc.add_paragraph('--------------------------------------')

# Function to add image to the Word document
def add_image_to_doc(doc, image_path, image_id):
    doc.add_paragraph(f"Image ID: {image_id}")
    doc.add_picture(image_path, width=Inches(2))
    doc.add_paragraph('--------------------------------------')

# Define global variables for camera and streaming state
cap = None
is_streaming = False

# Detection logic in a separate thread
def detection_loop():
    global cap, last_save_time, last_saved_boxes, image_counter, human_images
    while is_streaming:
        ret, frame = cap.read()
        if not ret:
            print("Failed to capture frame")
            break

        # Perform detection using YOLO
        results = model(frame)

        # Check if results contain any predictions
        if results[0].boxes:  # Access the boxes from results
            for box in results[0].boxes:  # Iterate through detected boxes
                x1, y1, x2, y2 = map(int, box.xyxy[0].tolist())  # Get bounding box coordinates
                conf = box.conf[0].item()  # Confidence score for detection
                cls = int(box.cls[0].item())  # Class index (0 for person, etc.)

                if conf > 0.5 and cls == 0:  # Class 0 corresponds to 'person'
                    current_time = time.time()

                    # Check if enough time has passed since the last save
                    if current_time - last_save_time > save_delay:
                        new_detection = True
                        for saved_box in last_saved_boxes:
                            if calculate_distance((x1, y1, x2, y2), saved_box) < proximity_threshold:
                                new_detection = False
                                break

                        if new_detection:
                            # Crop the detected person from the frame
                            cropped_person = frame[y1:y2, x1:x2]

                            # Save the cropped image
                            timestamp = time.strftime("%Y%m%d-%H%M%S")
                            image_filename = f"human_{timestamp}_{image_counter}.jpg"
                            image_path = os.path.join(output_folder, image_filename)
                            cv2.imwrite(image_path, cropped_person)
                            image_counter += 1  # Increment counter for the next image

                            # Store bounding box and update the last save time
                            last_saved_boxes.append((x1, y1, x2, y2))
                            last_save_time = current_time

                            # Add the saved image to the list of human detections
                            human_images.append((image_path, image_filename.split('.')[0]))  # Save image path and ID

                    # Draw the bounding box around the person
                    cv2.rectangle(frame, (x1, y1), (x2, y2), (0, 255, 0), 2)  # Green box
                    label = f"Person {conf:.2f}"  # Add confidence level as label
                    cv2.putText(frame, label, (x1, y1 - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (0, 255, 0), 2)

        # Show the frame with the bounding boxes
        cv2.imshow("Human Detection", frame)

        # Stop detection when 'q' is pressed
        if cv2.waitKey(1) & 0xFF == ord('q'):
            break

    cap.release()
    cv2.destroyAllWindows()

 # Function to start streaming
def start_streaming():
    global is_streaming, cap, human_images
    if not is_streaming:
        cap = cv2.VideoCapture(0)  # Reinitialize the camera if it's not streaming
        if not cap.isOpened():
            print("Error: Could not open video stream.")  # Debugging line to check video capture status
            return
        
        is_streaming = True
        print("Starting detection...")  # Debugging line to confirm starting detection
        human_images.clear()  # Clear the human_images list before starting a new recording
        threading.Thread(target=detection_loop, daemon=True).start()

# Function to stop streaming and generate report
def stop_streaming():
    global is_streaming
    if is_streaming:
        is_streaming = False
        print("Detection stopped. Generating report...")  # Debugging line to check if stop is working
        # Generate and display the report automatically
        generate_and_show_report()

# Function to generate and show the latest report
def generate_and_show_report():
    if not human_images:
        messagebox.showinfo("No Detections", "No detections were made to generate a report.")
        return

    # Create a new Word document
    doc = Document()
    create_doc_report(doc, len(human_images))

    for image_path, image_id in human_images:
        add_image_to_doc(doc, image_path, image_id)

    # Generate a unique filename for the report
    doc_filename = generate_unique_filename()

    # Save the document
    doc.save(doc_filename)
    print(f"Report saved to {doc_filename}")  # Debugging line to confirm saving

    # Open the saved report automatically
    os.startfile(doc_filename)  # Automatically open the report without asking for selection

# Function to display available reports in the folder (traditional dialog box)
def show_reports():
    # Open a file dialog to select a report
    report_path = filedialog.askopenfilename(
        initialdir=output_folder,
        title="Select a Report",
        filetypes=(("Word documents", "*.docx"), ("All files", "*.*"))
    )
    if report_path:
        os.startfile(report_path)  # Open the selected report

# Create the main UI window
root = Tk()
root.title("Human Detection System")
root.geometry("400x300")
root.config(bg="#2e2e2e")

# Add a title label
title_label = Label(root, text="Human Detection System", font=("Arial", 18, "bold"), fg="white", bg="#2e2e2e")
title_label.pack(pady=20)

# Start button
start_button = Button(root, text="Start Detection", width=20, font=("Arial", 12), bg="#4CAF50", fg="white", command=start_streaming)
start_button.pack(pady=10)

# Stop button
stop_button = Button(root, text="Stop Detection", width=20, font=("Arial", 12), bg="#f44336", fg="white", command=stop_streaming)
stop_button.pack(pady=10)

# Show Reports button
show_reports_button = Button(root, text="Show Reports", width=20, font=("Arial", 12), bg="#008CBA", fg="white", command=show_reports)
show_reports_button.pack(pady=10)

# Exit button
exit_button = Button(root, text="Exit", width=20, font=("Arial", 12), bg="#9E9E9E", fg="white", command=root.quit)
exit_button.pack(pady=10)

# Start the GUI main loop
root.mainloop()
